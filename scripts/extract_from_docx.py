#!/usr/bin/env python
"""
extract_from_docx.py — 从 Word 文档自动提取 ALS 患者数据

用法:
    # 读取 Word 文档，自动生成患者 JSON
    python scripts/extract_from_docx.py "患者报告.docx"

    # 生成 JSON 后直接跑预测
    python scripts/extract_from_docx.py "患者报告.docx" --predict

    # 指定输出文件名
    python scripts/extract_from_docx.py "患者报告.docx" -o patient.json

流程:
    Word文档 → 提取文本 → LLM解析 → 自动填充JSON → (可选)直接预测
"""

import sys
import json
import argparse
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from c9agent.utils.llm_client import run_llm


def extract_text_from_docx(filepath: str) -> str:
    """从 .docx 文件提取纯文本"""
    try:
        from docx import Document
    except ImportError:
        print("需要安装 python-docx: pip install python-docx")
        sys.exit(1)

    doc = Document(filepath)
    lines = []

    # 遍历段落
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            lines.append(text)

    # 遍历表格（常见于检验报告）
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                lines.append(" | ".join(cells))

    return "\n".join(lines)


def scan_xlsx_headers(filepath: str) -> dict:
    """
    扫描 Excel 表头，自动识别 ALS 相关列名。
    返回: {"alsfrsr_cols": [...], "fvc_col": "列名", ...}
    """
    import openpyxl
    wb = openpyxl.load_workbook(filepath, data_only=True)
    ws = wb.active

    # 收集所有单元格文本（前20行），用于关键词匹配
    all_text = ""
    for row in ws.iter_rows(values_only=True, max_row=20):
        cells = [str(c).strip() for c in row if c is not None]
        all_text += " ".join(cells) + "\n"

    hints = []

    # ALSFRS-R 相关关键词
    alsfrsr_keywords = [
        "ALSFRS", "ALS功能评分", "功能评分", "评分量表",
        "总分", "total_score", "total score", "frs",
        "修订版", "功能量表", "ALSFRS-R", "alsfrs",
    ]
    found_alsfrsr = [kw for kw in alsfrsr_keywords if kw.lower() in all_text.lower()]
    if found_alsfrsr:
        hints.append(f"[已识别] ALSFRS-R 相关列: {', '.join(found_alsfrsr)}")

    # FVC 关键词
    fvc_keywords = ["FVC", "用力肺活量", "肺活量", "forced vital capacity",
                    "肺功能", "呼吸功能", "%预计值", "%pred"]
    found_fvc = [kw for kw in fvc_keywords if kw.lower() in all_text.lower()]
    if found_fvc:
        hints.append(f"[已识别] FVC/呼吸功能: {', '.join(found_fvc)}")

    # 基因关键词
    gene_keywords = ["基因", "突变", "变异", "C9orf72", "SOD1", "TARDBP",
                     "FUS", "genetic", "gene", "variant", "检测"]
    found_gene = [kw for kw in gene_keywords if kw.lower() in all_text.lower()]
    if found_gene:
        hints.append(f"[已识别] 基因检测: {', '.join(found_gene)}")

    return hints


def extract_structured_from_xlsx(filepath: str) -> dict:
    """
    直接从 Excel 提取结构化字段（不经过 LLM）。
    返回已识别的字段，LLM 只需补全缺失的。

    策略: 遍历所有单元格，按关键词匹配 → 提取相邻单元格的值。
    """
    import openpyxl, re
    from datetime import datetime as dt

    wb = openpyxl.load_workbook(filepath, data_only=True)

    result = {
        "alsfrsr_records": [],
        "fvc_percent": None,
        "genetic_variants": [],
        "medications": [],
    }

    # 收集所有行（单元格文本 + 坐标）
    all_cells = []
    for sheet in wb:
        for row in sheet.iter_rows():
            for cell in row:
                if cell.value is not None:
                    all_cells.append({
                        "text": str(cell.value).strip(),
                        "row": cell.row,
                        "col": cell.column,
                    })

    # —— 1. 找 ALSFRS-R 总分 ——
    # 先找明确的总分行（"总分"、"/48"），再 fallback 到纯数字
    alsfrsr_total_pattern = re.compile(
        r'(ALSFRS.*总分|总分.*ALSFRS|total.*score|总分)',
        re.IGNORECASE
    )
    alsfrsr_any_pattern = re.compile(
        r'(ALSFRS|alsfrs|frs|功能评分)',
        re.IGNORECASE
    )
    score_in_text = re.compile(r'(\d{1,2})\s*/\s*48')  # "36/48" 格式
    score_simple = re.compile(r'^\s*(\d{1,2})\s*(?:分)?\s*$')

    total_score_val = None
    total_score_date = None

    # 第一遍: 优先找总分
    for cell in all_cells:
        if alsfrsr_total_pattern.search(cell["text"]):
            m = score_in_text.search(cell["text"])
            if m:
                total_score_val = int(m.group(1))
            else:
                # 同行找数字
                row, col = cell["row"], cell["col"]
                for other in all_cells:
                    if other["row"] == row and other["col"] > col:
                        m2 = score_simple.match(other["text"])
                        if m2:
                            val = int(m2.group(1))
                            if 0 <= val <= 48:
                                total_score_val = val
                                break
            break  # 找到总分就停

    # 第二遍: 如果没找到总分，从任意 ALSFRS 行取
    if total_score_val is None:
        for cell in all_cells:
            if alsfrsr_any_pattern.search(cell["text"]):
                m = score_in_text.search(cell["text"])
                if m:
                    total_score_val = int(m.group(1))
                    break
                row, col = cell["row"], cell["col"]
                for other in all_cells:
                    if other["row"] == row and other["col"] > col:
                        m2 = score_simple.match(other["text"])
                        if m2:
                            val = int(m2.group(1))
                            if 0 <= val <= 48:
                                total_score_val = val
                                break
                if total_score_val:
                    break

    # 收集同区域的日期
    date_found = None
    for cell in all_cells:
        for fmt in ["%Y-%m-%d", "%Y/%m/%d", "%Y.%m.%d", "%Y年%m月%d日"]:
            try:
                d = dt.strptime(cell["text"], fmt)
                date_found = d.strftime("%Y-%m-%d")
                # 找离 ALSFRS-R 最近的日期
                break
            except ValueError:
                continue

    # 也扫描"评估日期"、"检查日期"等关键词周围的日期
    for cell in all_cells:
        if any(kw in cell["text"] for kw in ["评估", "检查", "评定", "就诊", "入院", "记录"]):
            row, col = cell["row"], cell["col"]
            for other in all_cells:
                if other["row"] == row and other["col"] > col:
                    for fmt in ["%Y-%m-%d", "%Y/%m/%d", "%Y.%m.%d", "%Y年%m月%d日"]:
                        try:
                            d = dt.strptime(other["text"], fmt)
                            date_found = d.strftime("%Y-%m-%d")
                            break
                        except ValueError:
                            continue

    if total_score_val is not None:
        result["alsfrsr_records"].append({
            "date": date_found or "2026-08-01",
            "total_score": total_score_val,
        })

    # —— 2. 找 FVC ——
    fvc_pattern = re.compile(r'(FVC|fvc|用力肺活量|肺活量|%预计值|%pred)', re.IGNORECASE)
    pct_pattern = re.compile(r'(\d+)\s*%?')

    for i, cell in enumerate(all_cells):
        if fvc_pattern.search(cell["text"]):
            m = pct_pattern.search(cell["text"])
            if m:
                result["fvc_percent"] = int(m.group(1))
            else:
                # 检查同行相邻单元格
                row, col = cell["row"], cell["col"]
                for other in all_cells:
                    if other["row"] == row and other["col"] > col:
                        m2 = pct_pattern.search(other["text"])
                        if m2:
                            result["fvc_percent"] = int(m2.group(1))
                            break

    # —— 3. 找基因 ——
    gene_pattern = re.compile(
        r'(C9orf72|SOD1|TARDBP|FUS|TBK1|OPTN|VCP|UBQLN2|SQSTM1|'
        r'NEK1|KIF5A|ANXA11|CHCHD10|PFN1)',
        re.IGNORECASE
    )
    # 阳性/致病信号
    pathogenic_pattern = re.compile(r'(致病|pathogenic|阳性|\+|突变\s*$|变异\s*$)', re.IGNORECASE)
    # 阴性/排除信号 — 这些不算
    negative_pattern = re.compile(
        r'(野生型|阴性|排除|正常|未检出|未发现|无突变|wild\s*type|'
        r'benign|良性|多态性|polymorphism)',
        re.IGNORECASE
    )

    for cell in all_cells:
        m = gene_pattern.search(cell["text"])
        if m:
            gene = m.group(1).upper()
            # 收集该基因所在行的全部文本
            row = cell["row"]
            row_text = " ".join(
                c["text"] for c in all_cells if c["row"] == row
            )
            # 阴性结果跳过
            if negative_pattern.search(row_text):
                continue
            is_pathogenic = bool(pathogenic_pattern.search(row_text))
            if gene == "C9orf72":
                var_type = "repeat_expansion"
            else:
                var_type = "missense"

            result["genetic_variants"].append({
                "gene": gene,
                "variant_type": var_type,
                "acmg_classification": "P" if is_pathogenic else "VUS",
                "zygosity": "heterozygous",
            })

    # —— 4. 找用药 ——
    med_map = {
        "利鲁唑": "Riluzole", "力如太": "Riluzole", "riluzole": "Riluzole",
        "依达拉奉": "Edaravone", "edaravone": "Edaravone",
        "拉迪卡苷": "Radicava", "radicava": "Radicava",
        "托弗森": "Tofersen", "tofersen": "Tofersen",
    }
    for cell in all_cells:
        for zh, en in med_map.items():
            if zh.lower() in cell["text"].lower():
                if not any(m["drug_name"] == en for m in result["medications"]):
                    result["medications"].append({"drug_name": en, "dosage": None, "start_date": None})

    # 去重 ALSFRS-R（相同日期+分数）
    seen = set()
    unique_records = []
    for r in result["alsfrsr_records"]:
        key = (r["date"], r["total_score"])
        if key not in seen:
            seen.add(key)
            unique_records.append(r)
    result["alsfrsr_records"] = unique_records

    return result


def extract_text_from_xlsx(filepath: str) -> str:
    """从 .xlsx 文件提取纯文本（逐行读取所有单元格）"""
    try:
        import openpyxl
    except ImportError:
        print("需要安装 openpyxl: pip install openpyxl")
        sys.exit(1)

    wb = openpyxl.load_workbook(filepath, data_only=True)
    lines = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        lines.append(f"--- Sheet: {sheet_name} ---")
        for row in ws.iter_rows(values_only=True):
            # 过滤全空行
            cells = [str(c).strip() for c in row if c is not None]
            if cells:
                lines.append(" | ".join(cells))

    return "\n".join(lines)


def extract_text_from_txt(filepath: str) -> str:
    """从 .txt 文件读取"""
    with open(filepath, "r", encoding="utf-8") as f:
        return f.read()


def parse_with_llm(text: str, hints: list[str] = None) -> dict:
    """
    用 LLM 从文档文本中提取 ALS 临床字段。

    这是核心步骤 —— LLM 理解中文临床报告，
    自动找出年龄、起病、ALSFRS-R、基因等信息。
    """
    prompt = f"""你是一个 ALS 临床数据提取专家。从以下医疗文档中提取关键信息。

## 文档内容
```
{text[:3000]}
```
## 自动识别提示（优先使用这些列）
{chr(10).join(f'- {h}' for h in (hints or []))}

## 提取要求
请用 JSON 格式返回（只返回JSON，不要解释）。如果文档中没有某个信息，用 null。

**字段说明 & 映射规则:**

- sex: "male" 或 "female"
- age_at_onset: 发病年龄（数字）
- onset_site: "bulbar"(球部/延髓/构音/吞咽) / "spinal_cervical"(颈段/上肢) / "spinal_lumbar"(腰段/下肢) / "respiratory"(呼吸) / "unknown"
- diagnosis_date: 确诊日期 YYYY-MM-DD
- symptom_onset_date: 最早症状日期 YYYY-MM-DD
- alsfrsr_records: [{{"date":"YYYY-MM-DD","total_score":0-48}}]
- fvc_percent_predicted: FVC%数字
- genetic_variants: [{{"gene":"C9orf72/SOD1/TARDBP/FUS","acmg_classification":"P"}}]
- medications: [{{"drug_name":"Riluzole/Edaravone","start_date":"YYYY-MM-DD"}}]
- milestones: [{{"milestone_type":"diagnosis/wheelchair/gastrostomy/niv_dependency/invasive_ventilation/death","date":"YYYY-MM-DD"}}]
- clinical_notes: 100字摘要

返回JSON:
{{"patient_id":"","sex":"","age_at_onset":null,"onset_site":"unknown","diagnosis_date":"","symptom_onset_date":null,"alsfrsr_records":[],"fvc_percent_predicted":null,"genetic_variants":[],"medications":[],"milestones":[],"clinical_notes":""}}

规则: 不编造数据，中文映射: "利鲁唑/力如太"→"Riluzole", "依达拉奉"→"Edaravone", "球部/延髓"→"bulbar"
"""
    answer = run_llm(prompt, max_tokens=2000)
    try:
        # 剥离可能的 ```json ... ``` 包裹
        clean = answer
        if "```json" in clean:
            clean = clean.split("```json")[-1]
        if "```" in clean:
            clean = clean.split("```")[0]
        # 找 JSON 边界
        start = clean.find("{")
        end = clean.rfind("}") + 1
        if start == -1 or end <= start:
            raise ValueError("响应中无 JSON")
        # 尝试修复截断的 JSON（补上缺失的括号）
        json_str = clean[start:end]
        # 如果 JSON 不完整，尝试补全
        open_braces = json_str.count("{") - json_str.count("}")
        open_brackets = json_str.count("[") - json_str.count("]")
        json_str += "]" * open_brackets + "}" * open_braces
        return json.loads(json_str)
    except (json.JSONDecodeError, ValueError, IndexError) as e:
        print(f"[警告] LLM 解析失败: {e}")
        print(f"LLM 原始输出: {answer[:500]}...")
        return {"error": "LLM提取失败，请手动修改JSON", "raw_text": text[:1000]}


def build_patient_json(extracted: dict) -> dict:
    """
    将 LLM 提取的结果转换为系统可读的 ALSPatientData JSON。
    填充缺失字段的默认值，确保格式合法。
    """
    from datetime import date

    # 基本字段默认值 + 容错
    patient = {
        "patient_id": extracted.get("patient_id") or f"DOC-{date.today().strftime('%Y%m%d')}",
        "sex": extracted.get("sex") or "male",
        "age_at_onset": extracted.get("age_at_onset") or 55,
        "onset_site": extracted.get("onset_site", "unknown"),
        "diagnosis_date": _fix_date(extracted.get("diagnosis_date")) or date.today().isoformat(),
        "symptom_onset_date": _fix_date(extracted.get("symptom_onset_date")),
        "alsfrsr_records": _fix_alsfrsr_records(extracted.get("alsfrsr_records", [])),
        "respiratory": {
            "fvc_percent_predicted": extracted.get("fvc_percent_predicted"),
            "niv_usage_hours_per_day": extracted.get("niv_usage"),
        },
        "genetic_variants": extracted.get("genetic_variants", []),
        "medications": _fix_medications(extracted.get("medications", [])),
        "milestones": _fix_milestones(extracted.get("milestones", [])),
        "family_history_als": bool(extracted.get("family_history_als")),
        "clinical_notes": extracted.get("clinical_notes", ""),
    }

    # 确保至少有一条 milestone（确诊事件）
    if not any(m.get("milestone_type") == "diagnosis" for m in patient["milestones"]):
        patient["milestones"].insert(0, {
            "milestone_type": "diagnosis",
            "date": patient["diagnosis_date"],
        })

    # 修复 milestones 的日期缺失
    for m in patient["milestones"]:
        if not m.get("date"):
            m["date"] = patient["diagnosis_date"]

    # 修复 genetic_variants 的必填字段
    for v in patient["genetic_variants"]:
        if not v.get("variant_type"):
            v["variant_type"] = "unknown"
        if not v.get("zygosity"):
            v["zygosity"] = "heterozygous"

    return patient


# —— 容错修复函数 ——

def _fix_date(d: str) -> str | None:
    """修复不完整的日期：'2026-01' → '2026-01-01'"""
    if not d:
        return None
    d = str(d).strip()
    if len(d) == 7 and "-" in d:   # "2026-01"
        d += "-01"
    return d if len(d) >= 10 else None

def _fix_alsfrsr_records(records: list) -> list:
    """修复 ALSFRS-R 记录格式"""
    fixed = []
    for r in (records or []):
        if isinstance(r, dict):
            date_val = _fix_date(r.get("date")) or r.get("date")
            if not date_val:
                # 没有日期，用默认值（LLM/规则会尽量补）
                from datetime import date as dt_date
                date_val = dt_date.today().isoformat()
            fixed.append({
                "date": date_val,
                "total_score": int(r.get("total_score", 0)),
            })
    return fixed

def _fix_medications(meds: list) -> list:
    """修复用药记录格式，确保有 start_date"""
    fixed = []
    for m in (meds or []):
        if isinstance(m, dict) and m.get("drug_name"):
            entry = {"drug_name": m["drug_name"],
                     "dosage": m.get("dosage")}
            entry["start_date"] = _fix_date(m.get("start_date")) or "2025-01-01"
            fixed.append(entry)
    return fixed

def _fix_milestones(milestones: list) -> list:
    """修复里程碑日期格式，拆分合并的 milestone_type"""
    valid_types = {
        "diagnosis", "gastrostomy", "wheelchair", "loss_of_ambulation",
        "loss_of_speech", "loss_of_swallowing", "niv_dependency",
        "invasive_ventilation", "death",
    }
    fixed = []
    for m in (milestones or []):
        if not isinstance(m, dict):
            continue
        types = m.get("milestone_type", "")
        # 拆分 "diagnosis/wheelchair" → 两条记录
        for t in str(types).replace(" ", "").split("/"):
            t = t.strip()
            if t in valid_types:
                fixed.append({
                    "milestone_type": t,
                    "date": _fix_date(m.get("date")) or m.get("date", ""),
                })
    return fixed

def print_extraction_summary(extracted: dict, patient_json: dict):
    """打印提取摘要，让用户确认"""
    print("\n" + "=" * 60)
    print("LLM 自动提取结果")
    print("=" * 60)

    fields = [
        ("患者ID", patient_json["patient_id"]),
        ("性别", patient_json["sex"]),
        ("发病年龄", f"{patient_json['age_at_onset']}岁"),
        ("起病部位", patient_json["onset_site"]),
        ("确诊日期", patient_json["diagnosis_date"]),
        ("症状出现", patient_json.get("symptom_onset_date") or "未提及"),
        ("ALSFRS-R 记录", f"{len(patient_json['alsfrsr_records'])}次"),
        ("FVC%", f"{patient_json['respiratory']['fvc_percent_predicted']}%" if patient_json['respiratory']['fvc_percent_predicted'] else "未提及"),
        ("基因变异", " ".join(v['gene'] for v in patient_json['genetic_variants']) or "未提及"),
        ("用药", " ".join(m['drug_name'] for m in patient_json['medications']) or "未提及"),
        ("里程碑", f"{len(patient_json['milestones'])}个事件"),
        ("家族史", "是" if patient_json.get('family_history_als') else "否/未提及"),
    ]

    for label, value in fields:
        print(f"  {label:12s}: {value}")

    # 检查关键缺失
    missing = []
    if not patient_json["alsfrsr_records"]:
        missing.append("!! ALSFRS-R 评分缺失 — 这是预测最关键的指标")
    if not patient_json["respiratory"]["fvc_percent_predicted"]:
        missing.append("!! FVC 缺失 — 影响呼吸预后判断")
    if patient_json["onset_site"] == "unknown":
        missing.append("!! 起病部位未知 — 影响进展分型")

    if missing:
        print("\n[关键缺失] 手动补充以下字段可提高预测准确度:")
        for m in missing:
            print(f"  {m}")

    print(f"\n临床摘要: {extracted.get('clinical_notes', '无')[:200]}")
    print("=" * 60)


def main():
    parser = argparse.ArgumentParser(
        description="从 Word/TXT 文档提取 ALS 患者数据并生成 JSON"
    )
    parser.add_argument("file", type=str, help="Word (.docx) 或 TXT 文件路径")
    parser.add_argument("--output", "-o", type=str, default="patient.json",
                        help="输出 JSON 路径 (default: patient.json)")
    parser.add_argument("--predict", "-p", action="store_true",
                        help="提取后直接运行预测")
    parser.add_argument("--edit", "-e", action="store_true",
                        help="提取后打开 JSON 供手动编辑（Windows 记事本）")

    args = parser.parse_args()

    # —— Step 1: 提取文本 ——
    filepath = Path(args.file)
    if not filepath.exists():
        print(f"[错误] 文件不存在: {args.file}")
        sys.exit(1)

    ext = filepath.suffix.lower()
    structured = {}   # 规则提取的结构化字段（仅 xlsx）
    if ext == ".docx":
        print(f"提取 Word 文档: {filepath.name}")
        text = extract_text_from_docx(str(filepath))
        hints = []
    elif ext in (".xlsx", ".xls"):
        print(f"提取 Excel 文件: {filepath.name}")
        # 先用规则直接提取结构化字段
        structured = extract_structured_from_xlsx(str(filepath))
        if structured["alsfrsr_records"]:
            print(f"  规则提取 ALSFRS-R: {len(structured['alsfrsr_records'])}次记录")
        if structured["fvc_percent"]:
            print(f"  规则提取 FVC: {structured['fvc_percent']}%")
        if structured["genetic_variants"]:
            genes = [v['gene'] for v in structured['genetic_variants']]
            print(f"  规则提取 基因: {', '.join(genes)}")
        if structured["medications"]:
            meds = [m['drug_name'] for m in structured['medications']]
            print(f"  规则提取 用药: {', '.join(meds)}")
        hints = scan_xlsx_headers(str(filepath))
        text = extract_text_from_xlsx(str(filepath))
    elif ext == ".txt":
        print(f"读取文本文件: {filepath.name}")
        text = extract_text_from_txt(str(filepath))
        hints = []
    else:
        print(f"[错误] 不支持的文件格式: {ext}。仅支持 .docx / .xlsx / .txt")
        sys.exit(1)

    if not text.strip():
        print("[错误] 文档为空或无法读取内容")
        sys.exit(1)

    print(f"  提取到 {len(text)} 个字符")

    # —— Step 2: LLM 解析（补全规则提取不到的字段）——
    print("\nLLM 解析中（补全姓名/起病部位/日期等）...")
    extracted = parse_with_llm(text, hints)

    if "error" in extracted:
        # 对 Excel：如果规则已提取到关键数据，LLM 失败不致命
        if ext in (".xlsx", ".xls") and structured.get("alsfrsr_records"):
            print("[注意] LLM 解析失败，但规则已提取到 ALSFRS-R 等数据，继续处理")
            extracted = {}  # 用空字典，合并时全靠规则数据
        else:
            print(f"\n[失败] {extracted['error']}")
            with open(args.output, "w", encoding="utf-8") as f:
                json.dump(build_patient_json({}), f, ensure_ascii=False, indent=2)
            sys.exit(1)

    # —— Step 3: 合并规则提取结果（规则优先，无条件覆盖 LLM） ——
    # 注意：Excel 的规则提取比 LLM 更可靠，直接用规则结果
    if ext in (".xlsx", ".xls"):
        extracted["alsfrsr_records"] = structured.get("alsfrsr_records", [])
        extracted["fvc_percent_predicted"] = structured.get("fvc_percent")
        extracted["genetic_variants"] = structured.get("genetic_variants", [])
        extracted["medications"] = structured.get("medications", [])

    # —— Step 4: 构建 JSON ——
    patient_json = build_patient_json(extracted)
    print_extraction_summary(extracted, patient_json)

    # —— Step 4: 保存 ——
    with open(args.output, "w", encoding="utf-8") as f:
        json.dump(patient_json, f, ensure_ascii=False, indent=2, default=str)
    print(f"\n患者数据已保存: {args.output}")

    # —— Step 5: 可选编辑 ——
    if args.edit:
        import os
        os.system(f'notepad "{args.output}"')
        print("请在记事本中修改后按任意键继续...")
        input()
        # 重新加载编辑后的 JSON
        with open(args.output, "r", encoding="utf-8") as f:
            patient_json = json.load(f)

    # —— Step 6: 可选预测 ——
    if args.predict:
        print("\n" + "=" * 60)
        print("开始预测...")
        print("=" * 60)
        from c9agent.data.patient_schema import ALSPatientData, ALSPatientInput
        from c9agent.core.orchestrator import CentralOrchestrator

        try:
            patient = ALSPatientData(**patient_json)
        except Exception as e:
            print(f"[错误] JSON 格式有误: {e}")
            print("请检查并修正后重新运行:")
            print(f"  python scripts/run_single_patient.py --input {args.output}")
            sys.exit(1)

        input_data = ALSPatientInput(
            patient_id=patient.patient_id,
            structured_data=patient,
        )

        orch = CentralOrchestrator()
        report = orch.analyze_single_patient(input_data)

        report_path = args.output.replace(".json", "_report.json")
        md_path = args.output.replace(".json", "_report.md")
        with open(report_path, "w", encoding="utf-8") as f:
            json.dump(report, f, ensure_ascii=False, indent=2, default=str)
        with open(md_path, "w", encoding="utf-8") as f:
            f.write(orch.report_builder.to_markdown(report))

        print(f"\n预测报告已保存: {report_path}")
        print(f"可读报告已保存: {md_path}")

    # —— 提示下一步 ——
    if not args.predict:
        print(f"\n检查 JSON 无误后运行预测:")
        print(f"  python scripts/run_single_patient.py --input {args.output} -o result.json -m result.md")


if __name__ == "__main__":
    main()
